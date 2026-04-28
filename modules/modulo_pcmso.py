# =============================================================================
# MÓDULO PCMSO v7.1 — OCR automatico + AgenteMedicoIA
# Funções públicas:
#   extrair_texto_pdf, extrair_pgr_com_fallback, enriquecer_pgr_com_fispq,
#   processar_pcmso, gerar_html_pcmso, gerar_docx_rq61
# =============================================================================

import io
import os
import re
import unicodedata
from copy import deepcopy
from datetime import date

import pandas as pd

VERSAO_MODULO_PCMSO = "7.1 (OCR + AgenteMedicoIA)"

# ---------------------------------------------------------------------------
# Import do Agente Médico IA
# ---------------------------------------------------------------------------
try:
    from modules.agente_medico_ia import processar_cargo_ia
    _AGENTE_IA_DISPONIVEL = True
except ImportError:
    try:
        from agente_medico_ia import processar_cargo_ia
        _AGENTE_IA_DISPONIVEL = True
    except ImportError:
        _AGENTE_IA_DISPONIVEL = False

try:
    from data.dicionario_cas import DICIONARIO_CAS
except ImportError:
    DICIONARIO_CAS = {}


# ============================================================================
# 1 — EXTRACAO DE TEXTO DO PDF  (pdfplumber → PyMuPDF → OCR)
# ============================================================================

_MIN_CHARS_POR_PAGINA = 150  # menos que isso por página = provavel PDF protegido


def _texto_esta_vazio(texto: str, num_paginas: int) -> bool:
    """
    Retorna True se o texto extraido for insuficiente para o numero de paginas.
    Ex: PDF com 150 paginas, texto total < 22500 chars (150 * 150) → OCR necessario.
    """
    if not texto or not texto.strip():
        return True
    media_por_pagina = len(texto) / max(num_paginas, 1)
    return media_por_pagina < _MIN_CHARS_POR_PAGINA


def _extrair_ocr(data: bytes) -> str:
    """
    Fallback OCR: converte paginas do PDF em imagem e aplica pytesseract.
    Usa pre-processamento (escala de cinza + binarizacao Otsu) para melhorar
    a acuracia em documentos escaneados ou protegidos.
    """
    texto = ""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        from PIL import Image
        import numpy as np
        import cv2
    except ImportError as e:
        return f"[OCR indisponivel: {e}]"

    try:
        paginas = convert_from_bytes(data, dpi=250)
    except Exception as e:
        return f"[OCR falhou na conversao de paginas: {e}]"

    config_tess = "--psm 6 --oem 3"  # layout bloco unico, LSTM engine

    for i, img in enumerate(paginas):
        try:
            # Pre-processamento para melhorar OCR
            img_array = np.array(img.convert("RGB"))
            cinza = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            _, binaria = cv2.threshold(cinza, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            img_proc = Image.fromarray(binaria)

            t = pytesseract.image_to_string(img_proc, lang="por+eng", config=config_tess)
            if t and t.strip():
                texto += t + "\n"
        except Exception:
            # Fallback sem pre-processamento
            try:
                t = pytesseract.image_to_string(img, lang="por+eng", config=config_tess)
                texto += (t or "") + "\n"
            except Exception:
                pass

    return texto


def extrair_texto_pdf(pdf_file) -> str:
    """
    Extrai texto de um PDF com 3 camadas de fallback:
    1. pdfplumber
    2. PyMuPDF (fitz)
    3. OCR via pdf2image + pytesseract (ativa se texto < 150 chars/pagina)
    """
    if hasattr(pdf_file, "read"):
        pdf_file.seek(0)
        data = pdf_file.read()
    else:
        data = bytes(pdf_file)

    num_paginas = 1
    texto = ""

    # --- Camada 1: pdfplumber ---
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            num_paginas = max(len(pdf.pages), 1)
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texto += t + "\n"
        if not _texto_esta_vazio(texto, num_paginas):
            return texto
    except Exception:
        pass

    # --- Camada 2: PyMuPDF ---
    texto_fitz = ""
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        num_paginas = max(len(doc), 1)
        for page in doc:
            t = page.get_text()
            if t:
                texto_fitz += t + "\n"
        if not _texto_esta_vazio(texto_fitz, num_paginas):
            return texto_fitz
    except Exception:
        pass

    # --- Camada 3: OCR ---
    texto_ocr = _extrair_ocr(data)
    return texto_ocr if texto_ocr.strip() else (texto or texto_fitz or "")


# ============================================================================
# 2 — EXTRACAO / PARSER LOCAL DE PGR
# ============================================================================

def _normalizar(texto: str) -> str:
    if not texto:
        return ""
    nfkd = unicodedata.normalize("NFKD", str(texto))
    return nfkd.encode("ASCII", "ignore").decode("ASCII").lower().strip()


_RE_GHE = re.compile(
    r"(?i)^(?:GHE\s*[:\-\u2013]?\s*|GRUPO\s+HOMOG[E\u00ca]NEO\s+DE\s+EXPOSI[\u00c7C][\u00c3A]O\s*[:\-\u2013]?\s*)(.+)$"
)
_RE_CARGO = re.compile(
    r"(?i)^(?:CARGO\s+)(.+?)(?:\s*[-\u2013]\s*CBO[:\s]*\d+)?\s*$"
)
_RE_AGENTE = re.compile(
    r"(?i)(?:agente\s*(?:qu[i\u00ed]mico|f[i\u00ed]sico|biol[o\u00f3]gico|de\s+risco)?\s*[:\-\u2013]?\s*)(.+)$"
)


def _parsear_pgr_local(texto: str) -> list:
    linhas = texto.split("\n")
    blocos = []
    bloco_atual = None
    for linha in linhas:
        ls = linha.strip()
        if not ls:
            continue
        m_ghe = _RE_GHE.match(ls)
        m_cargo = _RE_CARGO.match(ls)
        if m_ghe:
            if bloco_atual:
                blocos.append(bloco_atual)
            bloco_atual = {"ghe": m_ghe.group(1).strip(), "cargos": [], "riscos_mapeados": []}
        elif m_cargo and bloco_atual is not None:
            nome = m_cargo.group(1).strip()
            if nome not in bloco_atual["cargos"]:
                bloco_atual["cargos"].append(nome)
        elif bloco_atual is not None:
            m_ag = _RE_AGENTE.match(ls)
            if m_ag:
                bloco_atual["riscos_mapeados"].append(
                    {"nome_agente": m_ag.group(1).strip(), "perigo_especifico": ""}
                )
    if bloco_atual:
        blocos.append(bloco_atual)
    return blocos


def extrair_pgr_com_fallback(texto_pgr: str):
    try:
        from parser_pgr import parsear_pgr_texto
        resultado = parsear_pgr_texto(texto_pgr)
        if resultado:
            return resultado, "local"
    except Exception:
        pass
    dados = _parsear_pgr_local(texto_pgr)
    fonte = "local" if dados else "vazio"
    return dados, fonte


# ============================================================================
# 3 — ENRIQUECIMENTO COM FISPQ
# ============================================================================

def enriquecer_pgr_com_fispq(dados_ghe: list, resultados_fispq: list) -> list:
    agentes = []
    for fispq in resultados_fispq:
        nome = fispq.get("nome_produto") or fispq.get("produto") or ""
        cas  = fispq.get("cas") or ""
        if nome:
            agentes.append({"nome_agente": nome, "perigo_especifico": cas})
        for comp in fispq.get("componentes", []):
            nc = comp.get("nome") or comp.get("substancia") or ""
            cc = comp.get("cas") or ""
            if nc:
                agentes.append({"nome_agente": nc, "perigo_especifico": cc})
    for ghe in dados_ghe:
        existentes = {r.get("nome_agente", "").lower() for r in ghe.get("riscos_mapeados", [])}
        for ag in agentes:
            if ag["nome_agente"].lower() not in existentes:
                ghe.setdefault("riscos_mapeados", []).append(ag)
    return dados_ghe


# ============================================================================
# 4 — MOTOR DE EXAMES
# ============================================================================

_EXAMES_MINIMOS_CANTEIRO = [
    {"nome": "Exame Cl\u00ednico",      "adm": True,  "per": "12", "mro": True,  "ret": True,  "dem": True},
    {"nome": "Audiometria",        "adm": True,  "per": "12", "mro": True,  "ret": False, "dem": True},
    {"nome": "Acuidade Visual",    "adm": True,  "per": "12", "mro": True,  "ret": False, "dem": False},
    {"nome": "Hemograma Completo", "adm": True,  "per": "12", "mro": True,  "ret": False, "dem": False},
    {"nome": "Glicemia em Jejum",  "adm": True,  "per": "12", "mro": True,  "ret": False, "dem": False},
    {"nome": "ECG",                "adm": True,  "per": "12", "mro": True,  "ret": False, "dem": False},
    {"nome": "Espirometria",       "adm": True,  "per": "24", "mro": True,  "ret": False, "dem": True},
    {"nome": "RX de T\u00f3rax OIT",   "adm": True,  "per": "60", "mro": True,  "ret": False, "dem": True},
]

_EXAMES_MINIMOS_ESCRIT = [
    {"nome": "Exame Cl\u00ednico", "adm": True, "per": "12", "mro": True, "ret": True, "dem": True},
]


def _bool_para_x(val) -> str:
    if isinstance(val, bool):
        return "X" if val else "-"
    if isinstance(val, str):
        return val.upper().strip() or "-"
    return "-"


def _per_para_str(per) -> str:
    try:
        return f"{int(per)}M"
    except (TypeError, ValueError):
        return str(per).upper().strip() if per else ""


def _riscos_para_lista_str(riscos_mapeados: list) -> list:
    resultado = []
    for r in riscos_mapeados:
        if isinstance(r, dict):
            nome = r.get("nome_agente") or r.get("nome") or ""
            perigo = r.get("perigo_especifico") or ""
            s = " ".join(filter(None, [nome, perigo])).strip()
            if s:
                resultado.append(s)
        elif isinstance(r, str) and r.strip():
            resultado.append(r.strip())
    return resultado


def _contexto_do_ghe(ghe_nome: str, riscos_str: list) -> dict:
    t = _normalizar(ghe_nome + " " + " ".join(riscos_str))
    return {
        "altura": any(x in t for x in ["altura", "nr-35", "nr35", "andaime", "cremalheira", "grua", "telhado"]),
        "confinado": any(x in t for x in ["confinado", "cisterna", "poco"]),
        "eletricidade": any(x in t for x in ["eletric", "nr-10", "nr10", "energizado", "choque"]),
        "maquinas_pesadas": any(x in t for x in ["maquina", "betoneira", "guindaste", "grua", "cremalheira"]),
    }


def _resolver_exames_cargo(cargo, riscos_str, contexto, e_canteiro):
    if _AGENTE_IA_DISPONIVEL:
        resultado = processar_cargo_ia(
            cargo=cargo,
            riscos=riscos_str,
            contexto=contexto,
            e_canteiro=e_canteiro,
        )
        return resultado.get("exames", []), resultado.get("chave_mestra", "")
    base = deepcopy(_EXAMES_MINIMOS_CANTEIRO if e_canteiro else _EXAMES_MINIMOS_ESCRIT)
    return base, None


# ============================================================================
# 5 — processar_pcmso  (assinatura original mantida)
# ============================================================================

def processar_pcmso(dados_ghe: list, tipo_ambiente: str = "canteiro") -> pd.DataFrame:
    linhas = []
    for ghe_item in dados_ghe:
        nome_ghe    = ghe_item.get("ghe") or ghe_item.get("nome_ghe") or "GHE sem nome"
        cargos      = ghe_item.get("cargos", [])
        riscos_mapeados = ghe_item.get("riscos_mapeados", [])
        riscos_str  = _riscos_para_lista_str(riscos_mapeados)
        exames_pre  = ghe_item.get("exames", [])

        if tipo_ambiente == "canteiro":
            e_canteiro = True
        elif tipo_ambiente == "escritorio":
            e_canteiro = False
        else:
            nome_n = _normalizar(nome_ghe)
            e_canteiro = not any(
                x in nome_n for x in ["escritorio", "administrativo", "engenharia", "planejamento", "gerencia"]
            )
            if "almoxarifado" in nome_n:
                e_canteiro = True

        contexto = _contexto_do_ghe(nome_ghe, riscos_str)

        for cargo in cargos:
            if exames_pre:
                if isinstance(exames_pre[0], dict):
                    exames_base = deepcopy(exames_pre)
                else:
                    exames_base = [
                        {"nome": str(e), "adm": True, "per": "12", "mro": True, "ret": False, "dem": False}
                        for e in exames_pre
                    ]
                if _AGENTE_IA_DISPONIVEL:
                    res_ia = processar_cargo_ia(cargo=cargo, riscos=riscos_str, contexto=contexto, e_canteiro=e_canteiro)
                    nomes_ok = {_normalizar(e.get("nome", "")) for e in exames_base}
                    for ex in res_ia.get("exames", []):
                        if _normalizar(ex.get("nome", "")) not in nomes_ok:
                            exames_base.append(ex)
                            nomes_ok.add(_normalizar(ex.get("nome", "")))
                    fonte = f"banco+agente_ia:{res_ia.get('chave_mestra', '')}"
                else:
                    fonte = "banco_pre_definido"
                exames_finais = exames_base
            else:
                exames_finais, chave = _resolver_exames_cargo(cargo, riscos_str, contexto, e_canteiro)
                fonte = f"agente_ia:{chave}" if _AGENTE_IA_DISPONIVEL else "fallback_minimo"

            for ex in exames_finais:
                nome_ex = ex.get("nome", "") if isinstance(ex, dict) else str(ex)
                adm = _bool_para_x(ex.get("adm", True) if isinstance(ex, dict) else True)
                per = _per_para_str(ex.get("per", "12") if isinstance(ex, dict) else "12")
                mro = _bool_para_x(ex.get("mro", True) if isinstance(ex, dict) else True)
                ret = _bool_para_x(ex.get("ret", False) if isinstance(ex, dict) else False)
                dem = _bool_para_x(ex.get("dem", False) if isinstance(ex, dict) else False)
                linhas.append({
                    "GHE / Setor": nome_ghe,
                    "Cargo": cargo,
                    "Exame": nome_ex,
                    "ADM": adm,
                    "PER": per,
                    "MRO": mro,
                    "RT": ret,
                    "DEM": dem,
                    "Justificativa": fonte,
                })

    cols = ["GHE / Setor", "Cargo", "Exame", "ADM", "PER", "MRO", "RT", "DEM", "Justificativa"]
    if not linhas:
        return pd.DataFrame(columns=cols)
    return pd.DataFrame(linhas)


# ============================================================================
# 6 — gerar_html_pcmso
# ============================================================================

def gerar_html_pcmso(df: pd.DataFrame, cabecalho: dict = None) -> str:
    if cabecalho is None:
        cabecalho = {}
    razao   = cabecalho.get("razao_social", "")
    cnpj    = cabecalho.get("cnpj", "")
    medico  = cabecalho.get("medico_rt", "")
    vig_ini = cabecalho.get("vig_ini", "")
    vig_fim = cabecalho.get("vig_fim", "")
    resp    = cabecalho.get("responsavel_tec", "")
    obra    = cabecalho.get("obra", "")
    hoje    = date.today().strftime("%d/%m/%Y")

    cs = "border:1px solid #ccc;padding:6px 8px;font-size:12px;"
    th = f"{cs}background:#084D22;color:white;text-align:center;"
    cols = ["GHE / Setor", "Cargo", "Exame", "ADM", "PER", "MRO", "RT", "DEM"]

    cab_html = f"""
    <div style="font-family:Arial,sans-serif;margin:0 auto;max-width:1100px;padding:20px;">
    <h2 style="color:#084D22;text-align:center;">PROGRAMA DE CONTROLE M\u00c9DICO DE SA\u00daDE OCUPACIONAL</h2>
    <h3 style="color:#084D22;text-align:center;">NR-07 \u2014 PCMSO</h3>
    <table style="width:100%;border-collapse:collapse;margin-bottom:20px;font-size:13px;">
      <tr><td style="padding:4px 8px;"><b>Empresa:</b> {razao}</td><td><b>CNPJ:</b> {cnpj}</td></tr>
      <tr><td><b>M\u00e9dico RT:</b> {medico}</td><td><b>Obra:</b> {obra}</td></tr>
      <tr><td><b>Vig\u00eancia:</b> {vig_ini} a {vig_fim}</td><td><b>Resp. SST:</b> {resp}</td></tr>
      <tr><td colspan="2"><b>Gerado em:</b> {hoje} \u2014 {VERSAO_MODULO_PCMSO}</td></tr>
    </table>
    <table style="width:100%;border-collapse:collapse;margin-bottom:30px;">
      <thead><tr>{''.join(f'<th style="{th}">{c}</th>' for c in cols)}</tr></thead>
      <tbody>
    """

    rows_html = ""
    for _, row in df.iterrows():
        rows_html += "<tr>" + "".join(
            f"<td style='{cs}'>{row.get(c, '') if c in df.columns else ''}</td>"
            for c in cols
        ) + "</tr>\n"

    rodape = f"""
      </tbody></table>
      <p style="font-size:11px;color:#888;text-align:center;margin-top:40px;">
        Gerado pelo Sistema Automa\u00e7\u00e3o SST \u2014 Seconci GO | {hoje}
      </p></div>
    """
    return f"<!DOCTYPE html><html><body>{cab_html}{rows_html}{rodape}</body></html>"


# ============================================================================
# 7 — gerar_docx_rq61
# ============================================================================

def gerar_docx_rq61(df: pd.DataFrame, cabecalho: dict = None) -> bytes:
    if cabecalho is None:
        cabecalho = {}
    try:
        from docx import Document
        from docx.shared import RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return df.to_csv(index=False).encode("utf-8")

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2)

    titulo = doc.add_heading("PROGRAMA DE CONTROLE M\u00c9DICO DE SA\u00daDE OCUPACIONAL", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if titulo.runs:
        titulo.runs[0].font.color.rgb = RGBColor(0x08, 0x4D, 0x22)

    doc.add_heading("NR-07 \u2014 PCMSO", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = [
        ("Empresa",      cabecalho.get("razao_social", "")),
        ("CNPJ",         cabecalho.get("cnpj", "")),
        ("M\u00e9dico RT",    cabecalho.get("medico_rt", "")),
        ("Obra/Unidade", cabecalho.get("obra", "")),
        ("Vig\u00eancia",     f"{cabecalho.get('vig_ini', '')} a {cabecalho.get('vig_fim', '')}"),
        ("Resp. SST",    cabecalho.get("responsavel_tec", "")),
        ("Gerado em",    date.today().strftime("%d/%m/%Y") + f" \u2014 {VERSAO_MODULO_PCMSO}"),
    ]
    t_meta = doc.add_table(rows=len(meta), cols=2)
    t_meta.style = "Table Grid"
    for i, (k, v) in enumerate(meta):
        t_meta.rows[i].cells[0].text = k
        t_meta.rows[i].cells[1].text = v
    doc.add_paragraph()

    colunas = ["GHE / Setor", "Cargo", "Exame", "ADM", "PER", "MRO", "RT", "DEM"]
    t = doc.add_table(rows=1, cols=len(colunas))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, col in enumerate(colunas):
        p = hdr[i].paragraphs[0]
        run = p.add_run(col)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "084D22")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, col in enumerate(colunas):
            cells[i].text = str(row.get(col, "")) if col in df.columns else ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
